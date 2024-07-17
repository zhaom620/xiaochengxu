# -*- coding: utf-8 -*-
"""
Created on Wed Jul 17 13:50:03 2024

@author: PC
"""

from flask import Flask, request, render_template, send_file
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Ensure upload and output directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

@app.route('/')
def upload_form():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file part"
    file = request.files['file']
    if file.filename == '':
        return "No selected file"
    if file:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        output_file = process_file(filepath)
        return send_file(output_file, as_attachment=True)

def process_file(filepath):
    try:
        df = pd.read_excel(filepath, dtype={'证件号码': str})
    except Exception as e:
        return f"读取Excel文件失败: {e}"

    if df.empty:
        return "Excel文件中没有数据"

    doc = Document()
    title = doc.add_heading('行政复议申请书', level=1)

    run = title.runs[0]
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(14)

    required_columns = ['姓名', '性别', '民族', '证件号码', '出生年月日', '住所']
    for column in required_columns:
        if column not in df.columns:
            return f"缺少必要的列: {column}"

    for index, row in df.iterrows():
        if row[required_columns].isnull().all():
            continue

        id_number = str(row['证件号码']).split('.')[0]
        content = (f"申请人{index + 1}：{row['姓名']}，{row['性别']}，{row['民族']}，"
                   f"身份证号：{id_number}，{row['出生年月日']}，地址：{row['住所']}")
        paragraph = doc.add_paragraph(content)
        
        for run in paragraph.runs:
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(14)

    output_file = os.path.join(app.config['OUTPUT_FOLDER'], 'output.docx')
    try:
        doc.save(output_file)
    except Exception as e:
        return f"保存Word文件失败: {e}"

    return output_file

if __name__ == '__main__':
    app.run(debug=True)
